import openai
from docx import Document
import requests
from io import BytesIO
import os

# -------- CONFIG --------
openai.api_key = os.getenv("OPENAI_API_KEY")  # Set your API key as ENV variable
IMAGE_SAVE_PATH = "question_image.png"
DOCX_OUTPUT_PATH = "Generated_Math_Assignment.docx"

# -------- BASE QUESTIONS --------
base_questions = """
1. Each student at Central Middle School wears a uniform consisting of 1 shirt and 1 pair of pants...
(Full table and question text)
2. The top view of a rectangular package of 6 tightly packed balls is shown...
(Full question text and image)
"""

# -------- LLM PROMPT --------
prompt = f"""
You are given the following base math questions:
{base_questions}

Your task:
- Create 2 similar math questions.
- Preserve LaTeX math formatting if present.
- If an image is needed, describe it clearly.
- Output in the following format for each question:

@title ...
@description ...
@question ...
@instruction ...
@difficulty ...
@Order ...
@option ...
@@option Correct Answer
@explanation ...
@subject ...
@unit ...
@topic ...
@plusmarks 1
"""

# -------- CALL LLM --------
response = openai.ChatCompletion.create(
    model="gpt-4o",
    messages=[{"role": "user", "content": prompt}],
    temperature=0.7
)

generated_text = response["choices"][0]["message"]["content"]

print("\nGenerated Questions:\n", generated_text)

# -------- IMAGE GENERATION (OPTIONAL) --------
image_prompt = "Top view of a rectangular box containing 8 tightly packed spheres, radius 3cm"
image_response = openai.Image.create(
    prompt=image_prompt,
    n=1,
    size="512x512"
)
image_url = image_response['data'][0]['url']

# Download the image
img_data = requests.get(image_url).content
with open(IMAGE_SAVE_PATH, 'wb') as f:
    f.write(img_data)

# -------- CREATE WORD DOC --------
doc = Document()
doc.add_heading("Generated Math Assessment", 0)
doc.add_paragraph(generated_text)

# Add image to doc
doc.add_picture(IMAGE_SAVE_PATH)

doc.save(DOCX_OUTPUT_PATH)

print(f"\n✅ Word document saved as {DOCX_OUTPUT_PATH}")
